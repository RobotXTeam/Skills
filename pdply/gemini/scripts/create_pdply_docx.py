#!/usr/bin/env python3
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Inches, Pt


BASE_DIR = Path.cwd()
MESHCORE_MODE = False
TOP_SECTION_COUNTER = 0


def add_paragraphs(doc, paragraphs):
    for text in paragraphs or []:
        add_body_paragraph(doc, str(text))


def format_paragraph(paragraph, before=6, after=6, line_spacing=1.2):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    format_paragraph(p, before=6, after=6)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    return p


def add_meshcore_heading(doc, title, level):
    global TOP_SECTION_COUNTER
    p = doc.add_paragraph()
    if level <= 2:
        TOP_SECTION_COUNTER += 1
        format_paragraph(p, before=15, after=6)
        prefix = f"{TOP_SECTION_COUNTER}. "
        r = p.add_run(prefix)
        r.font.name = "Arial"
        r.font.size = Pt(15)
        r.font.color.rgb = RGBColor(0x33, 0x70, 0xFF)
        r2 = p.add_run(str(title))
        r2.bold = True
        r2.font.name = "Arial"
        r2.font.size = Pt(15)
    elif level == 3:
        format_paragraph(p, before=13, after=6)
        r = p.add_run(str(title))
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(14)
    else:
        format_paragraph(p, before=8, after=6)
        r = p.add_run(str(title))
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(10)


def add_bullets(doc, bullets):
    for item in bullets or []:
        if isinstance(item, dict):
            p = doc.add_paragraph()
            format_paragraph(p, before=6, after=6)
            label = item.get("label")
            text = item.get("text", "")
            if label:
                run = p.add_run(str(label))
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(10.5)
                p.add_run(": " + str(text))
            else:
                p.add_run(str(text))
        else:
            add_body_paragraph(doc, str(item))


def add_numbered(doc, items):
    for item in items or []:
        add_body_paragraph(doc, str(item))


def add_tables(doc, tables):
    for table_data in tables or []:
        headers = table_data.get("headers") or []
        rows = table_data.get("rows") or []
        if not headers and rows:
            headers = ["" for _ in rows[0]]
        if not headers:
            continue
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.autofit = True
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = str(header)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
        for row in rows:
            cells = table.add_row().cells
            for i in range(len(headers)):
                cells[i].text = str(row[i]) if i < len(row) else ""
        doc.add_paragraph()


def resolve_image_path(path):
    p = Path(path)
    if p.is_absolute():
        return p
    candidate = BASE_DIR / p
    if candidate.exists():
        return candidate
    return Path.cwd() / p


def add_images(doc, images):
    for item in images or []:
        if isinstance(item, str):
            path = resolve_image_path(item)
            caption = ""
            width = 5.8
        else:
            path = resolve_image_path(item.get("path", ""))
            caption = item.get("caption", "")
            width = float(item.get("width_inches", 5.75 if MESHCORE_MODE else 5.8))
        if MESHCORE_MODE and width > 5.75:
            width = 5.75
        if not path.exists():
            p = doc.add_paragraph()
            p.add_run(f"[待补：图片未找到 {path}]").bold = True
            continue
        doc.add_picture(str(path), width=Inches(width))
        if caption:
            cap = add_body_paragraph(doc, str(caption))
            for run in cap.runs:
                run.italic = True
                run.font.size = Pt(9)


def add_section(doc, section):
    level = int(section.get("level", 2))
    title = section.get("title")
    if title:
        if MESHCORE_MODE:
            add_meshcore_heading(doc, title, level)
        else:
            doc.add_heading(str(title), level=max(1, min(level, 4)))
    add_paragraphs(doc, section.get("paragraphs"))
    add_bullets(doc, section.get("bullets"))
    add_numbered(doc, section.get("numbered"))
    add_tables(doc, section.get("tables"))
    add_images(doc, section.get("images"))
    for child in section.get("sections") or []:
        add_section(doc, child)


def set_default_font(doc):
    styles = doc.styles
    for style_name in ["Normal", "List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(10.5)


def set_no_margins_like_meshcore(doc):
    for section in doc.sections:
        section.left_margin = None
        section.right_margin = None
        section.top_margin = None
        section.bottom_margin = None


def main():
    if len(sys.argv) != 3:
        print("Usage: create_pdply_docx.py input.json output.docx", file=sys.stderr)
        return 2
    global BASE_DIR
    global MESHCORE_MODE
    global TOP_SECTION_COUNTER
    input_path = Path(sys.argv[1]).resolve()
    BASE_DIR = input_path.parent
    data = json.loads(input_path.read_text(encoding="utf-8"))
    MESHCORE_MODE = data.get("format") == "meshcore"
    TOP_SECTION_COUNTER = 0
    doc = Document()
    set_default_font(doc)
    if MESHCORE_MODE:
        set_no_margins_like_meshcore(doc)
    title = data.get("title")
    if title:
        if MESHCORE_MODE:
            p = doc.add_paragraph()
            format_paragraph(p, before=24, after=24)
            r = p.add_run(str(title))
            r.bold = True
            r.font.name = "Arial"
            r.font.size = Pt(26)
        else:
            doc.add_heading(str(title), level=1)
    for section in data.get("sections") or []:
        add_section(doc, section)
    out = Path(sys.argv[2])
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
